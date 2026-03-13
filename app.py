import os
from flask import Flask, render_template, request, session, redirect, url_for, flash
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user

from ats_analyzer import calculate_ats_score
from resume_parser import extract_resume_text
from skills_extractor import extract_skills
from job_matcher import match_jobs

from dotenv import load_dotenv
load_dotenv()

from docx import Document

app = Flask(__name__)
app.secret_key = "talentlens_secret_key"

# ================= DATABASE =================
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

# ================= USER MODEL =================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(200), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    resumes = db.relationship('Resume', backref='user', lazy=True)

# ================= RESUME MODEL =================
class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200))
    best_match = db.Column(db.String(100))
    skill_match_percentage = db.Column(db.Integer)
    skills = db.Column(db.Text)
    matched_skills = db.Column(db.Text)
    missing_skills = db.Column(db.Text)

    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ================= Upload folder =================
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ================= Extract DOCX =================
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        return "\n".join(full_text)
    except:
        return ""

# ================= ML Resume Strength =================
def calculate_resume_strength(resume_text):

    job_keywords = """
    python java sql data analysis machine learning
    communication teamwork leadership problem solving
    software development project management cloud docker
    """

    documents = [resume_text, job_keywords]

    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)

    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

    score = int(similarity * 100)

    # prevent very low score
    if score < 30:
        score += 25

    return score

# ================= HOME =================
@app.route('/', methods=['GET', 'POST'])
@login_required
def home():

    if request.method == 'POST':

        file = request.files.get('resume')

        if not file or file.filename == '':
            return "No file selected", 400

        if not allowed_file(file.filename):
            return "Invalid file type", 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        if filename.lower().endswith(".pdf"):
            resume_text = extract_resume_text(filepath)
        else:
            resume_text = extract_text_from_docx(filepath)

        resume_text = resume_text.lower()

        # Resume strength
        resume_strength_score = calculate_resume_strength(resume_text)

        if resume_strength_score <= 40:
            strength_label = "Weak Match"
        elif resume_strength_score <= 60:
            strength_label = "Moderate Match"
        elif resume_strength_score <= 80:
            strength_label = "Strong Match"
        else:
            strength_label = "Excellent Match"

        # Skill extraction
        extracted_data = extract_skills(resume_text)
        skills = extracted_data["skills"]

        ats_score, ats_breakdown = calculate_ats_score(resume_text, skills)

        results = match_jobs(skills)
        best_result = max(results, key=lambda x: x["score"])

        best_match = best_result["role"]
        skill_match_percentage = int(best_result["score"])
        matched_skills = ", ".join(best_result["matched_skills"])
        missing_skills = ", ".join(best_result["missing_skills"])

        # SAVE TO DATABASE
        new_resume = Resume(
            filename=filename,
            best_match=best_match,
            skill_match_percentage=skill_match_percentage,
            skills=", ".join(skills),
            matched_skills=matched_skills,
            missing_skills=missing_skills,
            user_id=current_user.id
        )

        db.session.add(new_resume)
        db.session.commit()

        return render_template(
            "result.html",
            skills=skills,
            best_match=best_match,
            skill_match_percentage=skill_match_percentage,
            ats_score=ats_score,
            ats_breakdown=ats_breakdown,
            resume_strength_score=resume_strength_score,
            strength_label=strength_label,
            matched_skills=matched_skills,
            missing_skills=missing_skills
        )

    return render_template("index.html")

# ================= COMPARE =================
@app.route('/compare')
@login_required
def compare():
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template("compare.html", resumes=resumes)

# ================= CONTACT =================
@app.route('/contact')
@login_required
def contact():
    return render_template("contact.html")
# ================= ABOUT =================
@app.route('/about')
@login_required
def about():
    return render_template("about.html")

# ================= AUTH =================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get("email")
        password = request.form.get("password")

        user = User.query.filter_by(email=email).first()

        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('home'))
        else:
            flash("Invalid email or password")

    return render_template("login.html")

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        email = request.form.get("email")
        password = generate_password_hash(request.form.get("password"))

        if User.query.filter_by(email=email).first():
            flash("Email already exists")
            return redirect(url_for('signup'))

        new_user = User(email=email, password=password)
        db.session.add(new_user)
        db.session.commit()

        flash("Account created! Please login.")
        return redirect(url_for('login'))

    return render_template("signup.html")

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ================= RUN =================
if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
    
